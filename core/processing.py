"""Модуль для обработки документов (чтение, конвертация)."""
import os
import shutil
import subprocess
from pathlib import Path
from typing import Optional, List, Callable
from docx import Document
import fitz


class DocumentProcessor:
    """Обработка документов: конвертация, чтение текста."""
    
    def __init__(self, log_callback: Optional[Callable[[str], None]] = None):
        self.libreoffice_path = self._find_libreoffice()
        self._log = log_callback or print

    def _find_libreoffice(self) -> Optional[str]:
        paths = [
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for path in paths:
            if Path(path).exists():
                return path
        return None

    def find_all_files(self, folder: str) -> List[Path]:
        """Найти все поддерживаемые файлы в папке рекурсивно."""
        files = []
        root = Path(folder)
        for ext in ['*.docx', '*.doc', '*.pdf']:
            files.extend(root.rglob(ext))
        return files

    def process_file(self, file_path: Path, output_dir: Path) -> Optional[Path]:
        """Конвертировать файл в docx если нужно."""
        suffix = file_path.suffix.lower()
        if suffix == '.docx':
            dest = output_dir / file_path.name
            shutil.copy2(file_path, dest)
            return dest
        elif suffix == '.doc':
            return self._convert_doc_to_docx(file_path, output_dir)
        elif suffix == '.pdf':
            return self._convert_pdf_to_docx(file_path, output_dir)
        return None

    def _convert_doc_to_docx(self, doc_path: Path, output_dir: Path) -> Optional[Path]:
        if not self.libreoffice_path:
            self._log(f"LibreOffice не найден, пропускаем {doc_path.name}")
            return None
        output_path = output_dir / doc_path.with_suffix('.docx').name
        try:
            subprocess.run([
                self.libreoffice_path,
                '--headless', '--convert-to', 'docx',
                '--outdir', str(output_dir), str(doc_path)
            ], check=True, capture_output=True, timeout=60)
            return output_path
        except Exception:
            return None

    def _convert_pdf_to_docx(self, pdf_path: Path, output_dir: Path) -> Optional[Path]:
        output_path = output_dir / pdf_path.with_suffix('.docx').name
        try:
            pdf = fitz.open(pdf_path)
            text = []
            for page_num in range(min(pdf.page_count, 30)):
                page = pdf[page_num]
                text.append(page.get_text())
            pdf.close()
            if text:
                doc = Document()
                doc.add_paragraph("\n".join(text))
                doc.save(output_path)
                return output_path
            return None
        except Exception:
            return None

    def read_text(self, docx_path: Path) -> str:
        """Читать текст из docx файла."""
        try:
            doc = Document(docx_path)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception:
            return ""
