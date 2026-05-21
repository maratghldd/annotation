"""Core modules for document analysis - multi-model pipeline implementation."""
import os
import time
import shutil
import urllib3
import requests
import subprocess
from pathlib import Path
from docx import Document
import fitz
from typing import List, Dict, Optional, Callable
from dataclasses import dataclass

# Отключаем предупреждения об отсутствии SSL-сертификата (как в эталоне)
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

@dataclass
class DocumentInfo:
    original_path: Path
    processed_path: Optional[Path] = None
    text_content: str = ""
    generated_title: str = ""
    status: str = "pending"
    error_message: str = ""

@dataclass
class AnalysisResult:
    file_name: str
    title: str
    status: str
    original_name: str = ""
    file_path: str = ""

class OllamaClient:
    def __init__(self, base_url: str = None):
        # Настройки из эталона пользователя
        raw_url = base_url or os.environ.get("OLLAMA_URL", "https://ollama.k2.iksi.edu")
        self.base_url = raw_url.rstrip("/")
        self.generate_url = f"{self.base_url}/api/generate"
        
        # Назначение моделей согласно заданию
        self.model_translate = "glm-4.7-flash:latest"  # Быстрая для перевода
        self.model_annotate = "qwen3:235b"            # Мощная для первой версии
        self.model_review = "deepseek-r1:32b"         # Рассуждающая для проверки
        
        self.verify = False # Как в эталоне

    def _call_ollama(self, model: str, prompt: str) -> str:
        payload = {
            "model": model,
            "prompt": prompt,
            "stream": False,
            "think": False,
        }
        try:
            response = requests.post(
                self.generate_url,
                json=payload,
                timeout=(10, 300),
                verify=self.verify,
            )
            if response.status_code == 200:
                return response.json().get("response", "").strip()
            return ""
        except Exception as e:
            print(f"Ошибка Ollama ({model}): {e}")
            return ""

    def translate_if_needed(self, text: str) -> str:
        """Переводит текст на русский, если он на украинском."""
        if not text or len(text.strip()) < 20:
            return text

        prompt = f"""Определи язык текста. Если текст на УКРАИНСКОМ языке, ПЕРЕВЕДИ его на РУССКИЙ. 
Если текст уже на РУССКОМ, просто верни его БЕЗ ИЗМЕНЕНИЙ.
Не добавляй никаких пояснений, только результат.

Текст:
{text[:2000]}"""
        
        translated = self._call_ollama(self.model_translate, prompt)
        return translated if translated else text

    def generate_initial_annotation(self, text: str, filename: str) -> str:
        """Создает первичную подробную аннотацию (заголовок)."""
        prompt = f"""Проанализируй текст документа и придумай развёрнутое, понятное название (8-15 слов) на РУССКОМ языке.
Название должно отражать суть документа.

Имя файла: {filename}
Текст:
{text[:3000]}

Название:"""
        return self._call_ollama(self.model_annotate, prompt)

    def review_and_finalize(self, text: str, initial_title: str) -> str:
        """Проверяет первичную аннотацию и выбирает самый достоверный вариант."""
        prompt = f"""Ты — эксперт-редактор. Проверь предложенный заголовок на соответствие тексту документа.
Если заголовок точный, верни его. Если он содержит ошибки или неточности, исправь его, сделав максимально достоверным.

Текст документа (отрывок):
{text[:2000]}

Предложенный заголовок:
{initial_title}

Итоговый, самый достоверный заголовок (только текст):"""
        
        final_title = self._call_ollama(self.model_review, prompt)
        return final_title if final_title else initial_title

    def check_connection(self) -> bool:
        try:
            requests.get(f"{self.base_url}/api/tags", timeout=5, verify=self.verify)
            return True
        except:
            return False

class DocumentProcessor:
    def __init__(self, log_callback: Optional[Callable[[str], None]] = None):
        self.libreoffice_path = self._find_libreoffice()
        self._log = log_callback or print

    def _find_libreoffice(self) -> Optional[str]:
        # Добавляем путь для macOS, так как пользователь на Mac
        paths = [
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
        ]
        for path in paths:
            if Path(path).exists():
                return path
        return None

    def find_all_files(self, folder: str) -> List[Path]:
        files = []
        root = Path(folder)
        for ext in ['*.docx', '*.doc', '*.pdf']:
            files.extend(root.rglob(ext))
        return files

    def process_file(self, file_path: Path, output_dir: Path) -> Optional[Path]:
        suffix = file_path.suffix.lower()
        if suffix == '.docx':
            dest = output_dir / file_path.name
            shutil.copy2(file_path, dest)
            return dest
        elif suffix == '.doc':
            return self._convert_doc_to_docx(file_path, output_dir)
        elif suffix == '.pdf':
            return self._convert_pdf_to_docx(file_path, output_dir)
        return None

    def _convert_doc_to_docx(self, doc_path: Path, output_dir: Path) -> Optional[Path]:
        if not self.libreoffice_path:
            self._log(f"LibreOffice не найден, пропускаем {doc_path.name}")
            return None
        output_path = output_dir / doc_path.with_suffix('.docx').name
        try:
            subprocess.run([
                self.libreoffice_path,
                '--headless', '--convert-to', 'docx',
                '--outdir', str(output_dir), str(doc_path)
            ], check=True, capture_output=True, timeout=60)
            return output_path
        except Exception:
            return None

    def _convert_pdf_to_docx(self, pdf_path: Path, output_dir: Path) -> Optional[Path]:
        output_path = output_dir / pdf_path.with_suffix('.docx').name
        try:
            pdf = fitz.open(pdf_path)
            text = []
            for page_num in range(min(pdf.page_count, 30)):
                page = pdf[page_num]
                text.append(page.get_text())
            pdf.close()
            if text:
                doc = Document()
                doc.add_paragraph("\n".join(text))
                doc.save(output_path)
                return output_path
            return None
        except Exception:
            return None

    def read_text(self, docx_path: Path) -> str:
        try:
            doc = Document(docx_path)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception:
            return ""

class ReportGenerator:
    @staticmethod
    def save_results(results: List[AnalysisResult], output_file: Path, source_folder: str):
        with open(output_file, "w", encoding="utf-8") as f:
            f.write("РЕЗУЛЬТАТЫ МНОГОЭТАПНОГО АНАЛИЗА\n")
            f.write(f"Источник: {source_folder}\n")
            f.write("=" * 60 + "\n\n")
            for i, r in enumerate(results, 1):
                f.write(f"[{i}] {r.original_name}\n   -> {r.title}\n" + "-"*40 + "\n")

class DocumentAnalyzer:
    def __init__(self, log_callback: Optional[Callable[[str], None]] = None, cancel_check: Optional[Callable[[], bool]] = None):
        self.ollama = OllamaClient()
        self.processor = DocumentProcessor(log_callback)
        self._log = log_callback or print
        self._cancel_check = cancel_check or (lambda: False)

    def run_folder_analysis(self, source: str, output: Path):
        self._log("Запуск многоэтапного анализа...")
        files = self.processor.find_all_files(source)
        results = []
        created_files = []

        for i, file_path in enumerate(files, 1):
            if self._cancel_check(): break
            self._log(f"[{i}/{len(files)}] Обработка: {file_path.name}")
            
            # Конвертация
            processed = self.processor.process_file(file_path, output)
            if not processed:
                results.append(AnalysisResult(file_path.name, "Ошибка обработки", "error", file_path.name))
                continue
            
            created_files.append(processed)
            text = self.processor.read_text(processed)
            
            # Этап 1: Перевод (если нужно)
            self._log("   - Этап 1: Проверка языка/Перевод (GLM-4 Flash)")
            working_text = self.ollama.translate_if_needed(text)
            
            # Этап 2: Первичная аннотация
            self._log("   - Этап 2: Первичная аннотация (Qwen 235B)")
            initial_title = self.ollama.generate_initial_annotation(working_text, file_path.name)
            
            # Этап 3: Проверка и финализация
            self._log("   - Этап 3: Проверка достоверности (DeepSeek R1 32B)")
            final_title = self.ollama.review_and_finalize(working_text, initial_title)
            
            results.append(AnalysisResult(processed.name, final_title, "success", file_path.name, str(processed)))
            self._log(f"   Готово: {final_title}")

        return results, created_files

    def run_single_file(self, file_path: Path, output_dir: Path) -> AnalysisResult:
        processed = self.processor.process_file(file_path, output_dir)
        if not processed:
            return AnalysisResult(file_path.name, "Ошибка", "error", file_path.name)
        
        text = self.processor.read_text(processed)
        self._log("Перевод...")
        text = self.ollama.translate_if_needed(text)
        self._log("Аннотирование...")
        initial = self.ollama.generate_initial_annotation(text, file_path.name)
        self._log("Проверка...")
        final = self.ollama.review_and_finalize(text, initial)
        
        return AnalysisResult(processed.name, final, "success", file_path.name, str(processed))